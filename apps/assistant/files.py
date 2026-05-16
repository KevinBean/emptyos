"""File-attachment text extraction — PDF, docx, txt, md.

Returns the extracted text plus a `truncated` flag so the assistant can warn
the user when the model only saw part of the document. PDF and docx parsers
are optional imports — when missing, the helper returns a clear error string
rather than crashing the chat turn.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

# Cap per-file extracted characters. 50_000 chars ≈ 12k tokens — fits inside
# every supported model's input window with room for chat history. Override
# via [apps.assistant] max_file_chars in emptyos.toml.
DEFAULT_MAX_CHARS = 50_000

EXTRACTABLE_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown"}


@dataclass
class ExtractedFile:
    path: str
    name: str
    text: str
    truncated: bool
    error: str = ""


def is_extractable_path(path: str) -> bool:
    if not path:
        return False
    return Path(path).suffix.lower() in EXTRACTABLE_EXTENSIONS


def extract_file(
    vault_root: str | Path, rel_path: str, *, max_chars: int = DEFAULT_MAX_CHARS
) -> ExtractedFile:
    abs_path = Path(vault_root) / rel_path
    name = abs_path.name
    if not abs_path.is_file():
        return ExtractedFile(rel_path, name, "", False, error="file not found")
    ext = abs_path.suffix.lower()
    try:
        if ext == ".pdf":
            text = _extract_pdf(abs_path)
        elif ext == ".docx":
            text = _extract_docx(abs_path)
        elif ext in (".txt", ".md", ".markdown"):
            text = abs_path.read_text(encoding="utf-8", errors="replace")
        else:
            return ExtractedFile(rel_path, name, "", False, error=f"unsupported file type: {ext}")
    except ImportError as e:
        return ExtractedFile(
            rel_path, name, "", False, error=f"parser missing — pip install required: {e}"
        )
    except Exception as e:
        return ExtractedFile(
            rel_path, name, "", False, error=f"{type(e).__name__}: {e}"
        )
    truncated = False
    if len(text) > max_chars:
        text = text[:max_chars]
        truncated = True
    return ExtractedFile(rel_path, name, text, truncated)


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except ImportError as e:
            raise ImportError("pypdf (or PyPDF2)") from e
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n\n".join(p for p in parts if p.strip())


def _extract_docx(path: Path) -> str:
    try:
        import docx  # type: ignore
    except ImportError as e:
        raise ImportError("python-docx") from e
    doc = docx.Document(str(path))
    lines: list[str] = []
    for para in doc.paragraphs:
        if para.text:
            lines.append(para.text)
    # Tables — flatten to tab-separated rows so the model sees structure
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append("\t".join(cells))
    return "\n".join(lines)


def format_block(extracted: ExtractedFile) -> str:
    """Render a single extracted file as a fenced markdown block for the LLM.

    The model is instructed to treat content between markers as context — keeps
    parsing predictable and lets the chat UI search/strip if it ever needs to.
    """
    if extracted.error:
        return (
            f"[Attached file `{extracted.name}` failed to read: {extracted.error}]"
        )
    suffix = " (truncated)" if extracted.truncated else ""
    return (
        f"[Attached file: {extracted.name}{suffix}]\n"
        f"```\n{extracted.text}\n```"
    )
