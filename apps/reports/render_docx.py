"""DOCX export — build a Word document programmatically from report source.

Fidelity is intentionally lower than PDF:
  - Headings, paragraphs, bold/italic/code, bullet/number lists, links, inline images, native tables.
  - Complex callouts and nested structures are flattened.
  - Figures are embedded at a sensible max width; captions are a separate paragraph.

python-docx must be installed (added to pyproject.toml as a reports-app runtime dep).
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from emptyos.sdk.utils import parse_frontmatter, strip_frontmatter

from . import tables as tables_mod
from .render_html import (
    FIGURE_RE,
    REQ_LINK_RE,
    TABLE_TOKEN_RE,
    WIKILINK_REQ_RE,
)


class PythonDocxMissing(RuntimeError):
    """Raised when python-docx isn't installed. Message includes the install command."""


INSTALL_HINT = (
    "DOCX export needs python-docx. Install with:\n"
    "    pip install python-docx"
)


def to_docx(report_dir: Path, out_path: Path) -> None:
    """Build a .docx from the report source in `report_dir` and save to `out_path`."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        raise PythonDocxMissing(INSTALL_HINT) from e

    meta_path = report_dir / "_meta.md"
    outline_path = report_dir / "_outline.md"
    sections_dir = report_dir / "sections"
    tables_dir = report_dir / "tables"
    figures_dir = report_dir / "figures"

    meta = _read_meta(meta_path)
    outline = _read_outline(outline_path)

    doc = Document()

    # --- Base style tweak ---
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
    except KeyError:
        pass

    # --- Title page ---
    _add_title_page(doc, meta, WD_ALIGN_PARAGRAPH, Pt)

    # --- TOC (placeholder — Word offers to update on open) ---
    doc.add_heading("Contents", level=1)
    toc_para = doc.add_paragraph()
    toc_para.add_run(
        "(Right-click and choose “Update Field” in Word to populate this table of contents.)"
    ).italic = True
    doc.add_page_break()

    # --- Sections ---
    figure_counter = {"n": 0}
    for section in outline:
        slug = section["slug"]
        title = section["title"]
        render_directive = section.get("render") or ""

        doc.add_heading(title, level=1)

        section_file = sections_dir / f"{slug}.md"
        body_md = ""
        if section_file.exists():
            raw = section_file.read_text(encoding="utf-8")
            body_md = strip_frontmatter(raw)

        # Strip/expand custom tokens inline
        body_md = _strip_table_tokens(body_md)
        body_md = _strip_req_links(body_md)
        figures = _collect_figures(body_md, figure_counter)
        body_md = FIGURE_RE.sub("", body_md)

        _render_markdown_into_doc(doc, body_md)

        for fig in figures:
            fig_path = figures_dir / fig["name"]
            if fig_path.exists():
                try:
                    doc.add_picture(str(fig_path), width=Inches(5.5))
                    cap = doc.add_paragraph(style="Caption" if "Caption" in [s.name for s in doc.styles] else None)
                    cap.add_run(f"Figure {fig['n']}: {fig['caption']}").italic = True
                except Exception:
                    # Skip bad images, continue.
                    p = doc.add_paragraph()
                    p.add_run(f"[Figure {fig['n']}: {fig['caption']} — could not embed {fig['name']}]").italic = True

        if render_directive.startswith("table:"):
            table_name = render_directive.split(":", 1)[1]
            rows = tables_mod.load_table(tables_dir / f"{table_name}.yaml")
            _add_structured_table(doc, table_name, rows)
        elif render_directive == "signoff":
            _add_signoff_table(doc, meta)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


# --- Helpers ---------------------------------------------------------------


def _read_meta(meta_path: Path) -> dict:
    if not meta_path.exists():
        return {}
    content = meta_path.read_text(encoding="utf-8")
    fm = parse_frontmatter(content) or {}
    return fm


def _read_outline(outline_path: Path) -> list[dict]:
    if not outline_path.exists():
        return []
    try:
        import yaml
    except ImportError:
        return []
    text = outline_path.read_text(encoding="utf-8")
    body = strip_frontmatter(text)
    body = re.sub(r"^#\s+Outline\s*\n", "", body, count=1)
    try:
        raw = yaml.safe_load(body)
    except Exception:
        return []
    if isinstance(raw, list):
        return [r for r in raw if isinstance(r, dict) and "slug" in r]
    return []


def _add_title_page(doc, meta: dict, WD_ALIGN_PARAGRAPH, Pt) -> None:
    doc_type = (meta.get("type") or "report").upper()
    title = meta.get("title") or "Untitled Report"
    subtitle = meta.get("subtitle") or ""
    version = meta.get("version") or "0.1"
    doc_date = meta.get("date") or date.today().isoformat()
    authors = meta.get("authors") or meta.get("author") or ""
    if isinstance(authors, list):
        authors = ", ".join(str(a) for a in authors)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(doc_type)
    r.bold = True
    r.font.size = Pt(14)

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if subtitle:
        sp = doc.add_paragraph(subtitle)
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1" if _style_exists(doc, "Light Grid Accent 1") else "Table Grid"
    meta_rows = [
        ("Document Type", doc_type),
        ("Version", str(version)),
        ("Date", str(doc_date)),
        ("Author(s)", str(authors)),
    ]
    org = meta.get("organisation") or ""
    if org:
        meta_rows.append(("Organisation", str(org)))
    project_id = meta.get("project_id") or ""
    if project_id:
        meta_rows.append(("Project", str(project_id)))
    for k, v in meta_rows:
        row = table.add_row()
        row.cells[0].text = k
        row.cells[1].text = v

    doc.add_page_break()


def _style_exists(doc, name: str) -> bool:
    try:
        return any(s.name == name for s in doc.styles)
    except Exception:
        return False


def _render_markdown_into_doc(doc, md_text: str) -> None:
    """Very small markdown-to-docx subset: headings, paragraphs, bullet/number lists, fenced code, bold/italic/inline code."""
    if not md_text or not md_text.strip():
        return

    lines = md_text.split("\n")
    in_code = False
    code_buf: list[str] = []
    para_buf: list[str] = []

    def flush_para():
        if not para_buf:
            return
        text = " ".join(l.strip() for l in para_buf if l.strip())
        if text:
            _add_inline_paragraph(doc, text)
        para_buf.clear()

    for line in lines:
        if line.strip().startswith("```"):
            if in_code:
                # close
                flush_code = "\n".join(code_buf)
                p = doc.add_paragraph()
                r = p.add_run(flush_code)
                r.font.name = "Consolas"
                code_buf.clear()
                in_code = False
            else:
                flush_para()
                in_code = True
            continue

        if in_code:
            code_buf.append(line)
            continue

        stripped = line.strip()

        if not stripped:
            flush_para()
            continue

        # Headings (## inside a section → level 2; ### → level 3)
        hm = re.match(r"^(#{2,4})\s+(.*)$", stripped)
        if hm:
            flush_para()
            level = min(len(hm.group(1)), 4)
            doc.add_heading(hm.group(2), level=level)
            continue

        # Bullet list
        if re.match(r"^[-*+]\s+", stripped):
            flush_para()
            text = re.sub(r"^[-*+]\s+", "", stripped)
            p = doc.add_paragraph(style="List Bullet") if _style_exists(doc, "List Bullet") else doc.add_paragraph()
            _apply_inline_runs(p, text)
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            flush_para()
            text = re.sub(r"^\d+\.\s+", "", stripped)
            p = doc.add_paragraph(style="List Number") if _style_exists(doc, "List Number") else doc.add_paragraph()
            _apply_inline_runs(p, text)
            continue

        para_buf.append(line)

    if in_code and code_buf:
        p = doc.add_paragraph()
        r = p.add_run("\n".join(code_buf))
        r.font.name = "Consolas"
    flush_para()


def _add_inline_paragraph(doc, text: str) -> None:
    p = doc.add_paragraph()
    _apply_inline_runs(p, text)


_INLINE_SPLIT_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def _apply_inline_runs(p, text: str) -> None:
    """Apply a minimal subset of inline markdown formatting to a paragraph."""
    parts = _INLINE_SPLIT_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 1:
            r = p.add_run(part[1:-1])
            r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Consolas"
        elif part.startswith("[") and "](" in part:
            label = part[1:part.index("](")]
            # python-docx lacks a first-class hyperlink API on paragraph level;
            # keep it simple and just add the label as italic (URL dropped in v1).
            r = p.add_run(label)
            r.italic = True
        else:
            p.add_run(part)


def _strip_table_tokens(body: str) -> str:
    """Remove {{table:x}} tokens — rendered separately as docx tables after the section prose."""
    return TABLE_TOKEN_RE.sub("", body)


def _strip_req_links(body: str) -> str:
    """Turn {{req:REQ-001}} / [[REQ-001]] into plain text 'REQ-001' in DOCX."""
    body = REQ_LINK_RE.sub(lambda m: m.group(2).upper(), body)
    body = WIKILINK_REQ_RE.sub(lambda m: m.group(1).upper(), body)
    return body


def _collect_figures(body: str, counter: dict) -> list[dict]:
    out: list[dict] = []
    for m in FIGURE_RE.finditer(body):
        counter["n"] += 1
        out.append({
            "n": counter["n"],
            "name": m.group(1).strip(),
            "caption": (m.group(3) or "").strip(),
        })
    return out


def _add_structured_table(doc, table_name: str, rows: list[dict]) -> None:
    from .templates import table_schema
    schema = table_schema(table_name)
    if schema is None:
        if not rows:
            p = doc.add_paragraph()
            p.add_run(f"No {table_name} recorded.").italic = True
            return
        columns = [{"key": k, "label": k.replace("_", " ").title()} for k in rows[0].keys()]
    else:
        columns = schema["columns"]

    if not rows:
        p = doc.add_paragraph()
        labels = " / ".join(c["label"] for c in columns)
        p.add_run(f"No {table_name} recorded — columns: {labels}.").italic = True
        return

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Light Grid Accent 1" if _style_exists(doc, "Light Grid Accent 1") else "Table Grid"
    hdr = table.rows[0].cells
    for i, c in enumerate(columns):
        hdr[i].text = c["label"]
    for r in rows:
        row = table.add_row()
        for i, c in enumerate(columns):
            val = r.get(c["key"], "")
            if isinstance(val, list):
                val = ", ".join(str(v) for v in val)
            row.cells[i].text = str(val)


def _add_signoff_table(doc, meta: dict) -> None:
    approvers = meta.get("approvers") or []
    if isinstance(approvers, str):
        approvers = [a.strip() for a in approvers.split(",") if a.strip()]
    if not approvers:
        p = doc.add_paragraph()
        p.add_run("No approvers configured.").italic = True
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1" if _style_exists(doc, "Light Grid Accent 1") else "Table Grid"
    h = table.rows[0].cells
    h[0].text, h[1].text, h[2].text, h[3].text = "Role", "Name", "Date", "Signature"
    for a in approvers:
        row = table.add_row()
        if isinstance(a, dict):
            row.cells[0].text = str(a.get("role", ""))
            row.cells[1].text = str(a.get("name", ""))
            row.cells[2].text = str(a.get("date", ""))
            row.cells[3].text = str(a.get("signature", ""))
        else:
            row.cells[0].text = str(a)
